from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SERVICES_PATH = ROOT / "content" / "services" / "services.json"
OUTPUT_PATH = ROOT / "client-documents" / "Advanced-Carpets-Service-Content-Review.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "E8EEF5"
FORM_FILL = "F4F6F9"
BORDER = "B8C2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
      shd = OxmlElement("w:shd")
      tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def set_run_font(run, size: int | float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str = "", style: str | None = None, bold: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, bold=True, color=BLUE if level < 3 else DARK_BLUE)


def add_field_block(doc: Document, title: str, current: list[str], blank_lines: int = 5) -> None:
    add_heading(doc, title, 3)
    add_para(doc, "Current draft", bold=True)
    for item in current:
        if item.strip():
            add_para(doc, item)

    table = doc.add_table(rows=2, cols=1)
    set_table_geometry(table, [6.5])
    header = table.rows[0].cells[0]
    body = table.rows[1].cells[0]
    set_cell_shading(header, LIGHT_FILL)
    set_cell_shading(body, FORM_FILL)
    header.paragraphs[0].add_run("Client changes / missing information")
    set_run_font(header.paragraphs[0].runs[0], bold=True, color=DARK_BLUE)
    body.paragraphs[0].add_run("\n".join([""] * blank_lines))


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1.75, 4.75])
    for idx, (label, value) in enumerate(rows):
        if idx:
            table.add_row()
        left, right = table.rows[idx].cells
        set_cell_shading(left, LIGHT_FILL)
        left.text = label
        right.text = value
        for paragraph in left.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True, color=DARK_BLUE)
        for paragraph in right.paragraphs:
            for run in paragraph.runs:
                set_run_font(run)


def format_sections(service: dict) -> list[str]:
    lines: list[str] = []
    for section in service.get("sections", []):
        lines.append(section["heading"])
        for paragraph in section.get("paragraphs", []) or []:
            lines.append(paragraph)
        for bullet in section.get("bullets", []) or []:
            lines.append(f"- {bullet}")
    return lines or ["No additional detail sections supplied yet."]


def format_faqs(service: dict) -> list[str]:
    return [f"Q: {faq['question']}\nA: {faq['answer']}" for faq in service.get("faqs", [])]


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Advanced Carpets service content review")
    set_run_font(run, size=9, color="666666")
    return doc


def add_service(doc: Document, service: dict, index: int) -> None:
    if index:
        doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, service["name"], 1)
    add_para(doc, "Use this section to approve the current draft, correct anything that is inaccurate, and add missing information the website should include.")

    add_key_value_table(
        doc,
        [
            ("Service name", service["name"]),
            ("URL path", f"/services/{service['slug']}/"),
            ("Category", service["category"]),
            ("Draft SEO title", f"{service['name']} | Advanced Carpets & Restoration"),
            ("Draft meta description", service["summary"]),
        ],
    )

    add_field_block(
        doc,
        "Approval and priority notes",
        [
            "[ ] Approved as written",
            "[ ] Approved with edits below",
            "[ ] Please replace this section",
            "Priority notes: accuracy, missing details, preferred wording, or services not offered.",
        ],
        blank_lines=6,
    )
    add_field_block(doc, "Hero intro", [service["intro"]], blank_lines=5)
    add_field_block(doc, "Overview and body copy", service.get("overviewParagraphs", []), blank_lines=7)
    add_field_block(doc, "Detailed service sections", format_sections(service), blank_lines=8)
    add_field_block(doc, "Benefits / reasons to book", service.get("benefits", []), blank_lines=6)
    add_field_block(doc, "Suitable for / use cases", service.get("useCases", []), blank_lines=6)
    add_field_block(doc, "Process steps", service.get("process", []), blank_lines=6)
    add_field_block(doc, "FAQs", format_faqs(service), blank_lines=8)
    add_field_block(
        doc,
        "Images, captions, and alt text",
        [
            f"Current caption: {service.get('imageCaption', 'No caption supplied yet.')}",
            f"Current results copy: {service.get('resultsCopy', 'No results copy supplied yet.')}",
            f"Suggested alt text: {service['name']} work in progress by Advanced Carpets & Restoration.",
            "Please list real job photos, before/after images, or videos that should be used for this service.",
        ],
        blank_lines=8,
    )
    add_field_block(
        doc,
        "Pricing, timing, qualifications, guarantees, and exclusions",
        [
            f"Current why-us copy: {service.get('whyAdvanced', 'No service-specific why-us copy supplied yet.')}",
            "Please add any pricing guidance, minimum job details, service area limits, qualifications, expected timing, drying time, safety notes, guarantees, or exclusions the website should mention.",
        ],
        blank_lines=8,
    )


def main() -> None:
    services = json.loads(SERVICES_PATH.read_text())["services"]
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = setup_document()
    title = doc.add_paragraph()
    title_run = title.add_run("Advanced Carpets Service Content Review")
    set_run_font(title_run, size=22, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Client review form for service page content, SEO copy, imagery, and missing details.")
    set_run_font(subtitle_run, size=12, color="555555")
    add_para(
        doc,
        "How to use this document: review each service section, tick the approval line that fits, and type corrections or missing information into the response boxes. The current draft text is prefilled from the website and the supplied service PDFs where available.",
    )
    add_bullets(
        doc,
        [
            "Keep wording practical and specific to the real service.",
            "Add real timing, access, safety, qualification, and exclusion details where customers should know them before enquiring.",
            "List any photos or videos that should replace the current website examples.",
        ],
    )

    for index, service in enumerate(services):
        add_service(doc, service, index)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
